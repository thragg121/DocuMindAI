from pathlib import Path

import fitz
from docx import Document as DocxDocument


class DocumentParsingError(Exception):
    """Raised when document text extraction fails."""


class DocumentParser:
    @staticmethod
    def extract_text(file_path: Path) -> str:
        extension = file_path.suffix.lower()

        try:
            if extension == ".pdf":
                text = DocumentParser._extract_pdf(file_path)
            elif extension == ".docx":
                text = DocumentParser._extract_docx(file_path)
            elif extension == ".txt":
                text = DocumentParser._extract_txt(file_path)
            else:
                raise DocumentParsingError(f"Unsupported document type: {extension}")
        except DocumentParsingError:
            raise
        except Exception as error:
            raise DocumentParsingError(
                f"Failed to extract text from " f"'{file_path.name}': {error}"
            ) from error

        cleaned_text = DocumentParser._clean_text(text)

        if not cleaned_text:
            raise DocumentParsingError("No readable text was found in the document.")

        return cleaned_text

    @staticmethod
    def _extract_pdf(file_path: Path) -> str:
        pages: list[str] = []

        with fitz.open(str(file_path)) as pdf_document:
            for page in pdf_document:
                page_text = page.get_text("text")

                if page_text.strip():
                    pages.append(page_text)

        return "\n\n".join(pages)

    @staticmethod
    def _extract_docx(file_path: Path) -> str:
        document = DocxDocument(str(file_path))

        paragraphs = [
            paragraph.text
            for paragraph in document.paragraphs
            if paragraph.text.strip()
        ]

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]

                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    @staticmethod
    def _extract_txt(file_path: Path) -> str:
        encodings = (
            "utf-8",
            "utf-8-sig",
            "cp1251",
            "latin-1",
        )

        for encoding in encodings:
            try:
                return file_path.read_text(
                    encoding=encoding,
                )
            except UnicodeDecodeError:
                continue

        raise DocumentParsingError("Unable to determine the TXT file encoding.")

    @staticmethod
    def _clean_text(text: str) -> str:
        lines = [line.strip() for line in text.replace("\x00", "").splitlines()]

        cleaned_lines: list[str] = []
        previous_line_was_empty = False

        for line in lines:
            if line:
                cleaned_lines.append(line)
                previous_line_was_empty = False
            elif not previous_line_was_empty:
                cleaned_lines.append("")
                previous_line_was_empty = True

        return "\n".join(cleaned_lines).strip()
